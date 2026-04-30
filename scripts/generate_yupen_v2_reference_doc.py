from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOCX = ROOT / "docs" / "yupen-v2-reference-validation-2026-03-31.docx"
OUTPUT_MD = ROOT / "analysis" / "yupen-v2-reference-validation-2026-03-31.md"


@dataclass
class Row:
    date: str
    strength: int
    code: str
    name: str
    status: str
    change_pct: Optional[float]
    current: Optional[float]
    threshold: Optional[float]
    provided_deviation: Optional[float]
    volume_ratio: Optional[float]
    turn_date: Optional[str]
    interval_pct: Optional[float]
    note: str = ""

    @property
    def calculated_deviation(self) -> Optional[float]:
        if self.current is None or self.threshold in (None, 0):
            return None
        return round((self.current / self.threshold - 1) * 100, 2)

    @property
    def deviation_match(self) -> Optional[bool]:
        if self.provided_deviation is None or self.calculated_deviation is None:
            return None
        return abs(self.provided_deviation - self.calculated_deviation) <= 0.03

    @property
    def status_match(self) -> Optional[bool]:
        if self.current is None or self.threshold is None:
            return None
        expected = "YES" if self.current >= self.threshold else "NO"
        return expected == self.status


def fmt_number(value: Optional[float]) -> str:
    if value is None:
        return "-"
    if abs(value) >= 1000 and float(value).is_integer():
        return f"{int(value)}"
    if float(value).is_integer():
        return f"{int(value)}"
    return f"{value:.2f}"


def fmt_pct(value: Optional[float]) -> str:
    if value is None:
        return "-"
    sign = "+" if value > 0 else ""
    return f"{sign}{value:.2f}%"


def set_cn_font(run, font_name: str = "PingFang SC", size: int = 10) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


DATA: list[Row] = [
    Row("2026-03-31", 1, "SPY", "标普500", "NO", 2.91, 650, 660, -1.52, 1.57, "26.02.27", -5.11),
    Row("2026-03-31", 2, "HSI", "恒生指数", "NO", 0.15, 24788, 25370, -2.29, 1.08, "26.02.24", -6.78),
    Row("2026-03-31", 3, "QQQ", "纳指100", "NO", 3.41, 577, 591, -2.37, 1.36, "26.03.12", -3.35),
    Row("2026-03-31", 4, "883418", "微盘股", "NO", -1.69, 2170, 2223, -2.38, 0.68, "26.03.19", -2.65),
    Row("2026-03-31", 5, "HSCEI", "国企指数", "NO", -0.30, 8374, 8584, -2.45, 0.95, "26.03.19", -3.69),
    Row("2026-03-31", 6, "399006", "创业板指", "NO", -2.70, 3184, 3278, -2.87, 0.96, "26.03.30", -2.72),
    Row("2026-03-31", 7, "AUUSDO", "黄金现价", "NO", 3.89, 4688, 4828, -2.90, None, "26.03.12", -7.99),
    Row("2026-03-31", 8, "399300", "沪深300", "NO", -0.93, 4450, 4586, -2.97, 0.99, "26.03.13", -4.69),
    Row("2026-03-31", 9, "1B0016", "上证50", "NO", -0.25, 2826, 2914, -3.02, 1.04, "26.02.13", -8.28),
    Row("2026-03-31", 10, "AGUSDO", "白银现价", "NO", 7.28, 75.00, 77.42, -3.13, None, "26.03.13", -6.90),
    Row("2026-03-31", 11, "000510", "中证A500", "NO", -1.20, 5526, 5725, -3.48, 0.97, "26.03.13", -5.44),
    Row("2026-03-31", 12, "932000", "中证2000", "NO", -1.67, 3242, 3382, -4.14, 0.99, "26.03.13", -6.60),
    Row("2026-03-31", 13, "1B0852", "中证1000", "NO", -1.91, 7619, 7976, -4.48, 0.92, "26.03.13", -7.24),
    Row("2026-03-31", 14, "399905", "中证500", "NO", -1.76, 7617, 8005, -4.85, 0.94, "26.03.12", -8.88),
    Row("2026-03-31", 15, "N225", "日经225", "NO", -1.58, 51063, 53763, -5.02, None, "26.03.03", -9.27, "临界值按偏离率反推修正，原始 OCR 文本疑似缺位。"),
    Row("2026-03-31", 16, "HS2083", "恒生科技", "NO", -0.86, 4908, 5183, -5.28, 0.77, "26.03.19", -6.95),
    Row("2026-03-31", 17, "1B0688", "科创50", "NO", -2.59, 1347, 1570, -14.21, 0.97, "26.03.02", -14.21, "临界值按偏离率反推取整，原始 OCR 文本错位。"),
    Row("2026-03-31", 18, "899050", "北证50", "NO", -0.99, 1248, 1347, -7.35, 1.02, "26.02.27", -18.80),
    Row("2026-03-30", 1, "399006", "创业板指", "NO", -0.68, 3273, 3280, -0.21, 0.86, "26.03.30", -0.67),
    Row("2026-03-30", 2, "883418", "微盘股", "NO", 0.41, 2207, 2224, -0.76, 0.87, "26.03.19", -0.99),
    Row("2026-03-30", 3, "399300", "沪深300", "NO", -0.24, 4491, 4596, -2.28, 0.84, "26.03.13", -3.81),
    Row("2026-03-30", 4, "HSCEI", "国企指数", "NO", -0.65, 8399, 8596, -2.29, 0.97, "26.03.19", -3.40),
    Row("2026-03-30", 5, "000510", "中证A500", "NO", -0.11, 5594, 5740, -2.54, 0.87, "26.03.13", -4.28),
    Row("2026-03-30", 6, "HSI", "恒生指数", "NO", -0.81, 24750, 25419, -2.63, 0.96, "26.02.24", -6.92),
    Row("2026-03-30", 7, "932000", "中证2000", "NO", 0.37, 3297, 3391, -2.77, 0.93, "26.03.13", -5.01),
    Row("2026-03-30", 8, "1B0852", "中证1000", "NO", 0.28, 7767, 8002, -2.94, 0.88, "26.03.13", -5.44),
    Row("2026-03-30", 9, "1B0016", "上证50", "NO", -0.14, 2833, 2924, -3.11, 0.78, "26.02.13", -8.05),
    Row("2026-03-30", 10, "399905", "中证500", "NO", 0.21, 7753, 8038, -3.55, 0.87, "26.03.12", -7.25),
    Row("2026-03-30", 11, "N225", "日经225", "NO", -2.79, 51885, 54109, -4.11, None, "26.03.03", -7.81),
    Row("2026-03-30", 12, "HS2083", "恒生科技", "NO", -1.84, 4690, 4920, -4.67, 0.78, "26.03.19", -6.12),
    Row("2026-03-30", 13, "SPY", "标普500", "NO", -0.33, 631, 662, -4.68, 0.95, "26.02.27", -7.88),
    Row("2026-03-30", 14, "1B0688", "科创50", "NO", -0.84, 1289, 1353, -4.73, 0.86, "26.03.02", -11.95),
    Row("2026-03-30", 15, "QQQ", "纳指100", "NO", -0.76, 558, 592, -5.74, 0.95, "26.03.12", -6.53),
    Row("2026-03-30", 16, "AUUSDO", "黄金现价", "NO", 1.48, 4850, 4513, None, 0.83, "26.02.25", 18.02, "原始文本内部冲突: 现价高于临界值却标注 NO，本行仅保留原文参考，不纳入状态校验。"),
    Row("2026-03-30", 17, "899050", "北证50", "NO", -0.84, 1260, 1355, -7.01, None, None, None, "量比/区间涨幅 OCR 缺失。"),
    Row("2026-03-30", 18, "AGUSDO", "白银现价", "NO", 2.56, 69.83, 77.82, -10.27, None, "26.03.13", -13.32),
    Row("2026-03-26", 1, "399006", "创业板指", "NO", -1.34, 3272, 3281, -0.27, 0.86, "26.03.26", -1.33),
    Row("2026-03-26", 2, "N225", "日经225", "NO", -0.27, 53603, 54727, -2.05, None, "26.03.03", -4.75),
    Row("2026-03-26", 3, "HSI", "恒生指数", "NO", -1.89, 24856, 25569, -2.79, 0.83, "26.02.24", -6.52),
    Row("2026-03-26", 4, "HSCEI", "国企指数", "NO", -2.25, 8389, 8631, -2.80, 0.95, "26.03.19", -3.52),
    Row("2026-03-26", 5, "399300", "沪深300", "NO", -1.32, 4477, 4619, -3.07, 0.79, "26.03.13", -4.11),
    Row("2026-03-26", 6, "883418", "微盘股", "NO", -1.58, 2161, 2234, -3.27, 0.77, "26.03.19", -3.05),
    Row("2026-03-26", 7, "SPY", "标普500", "NO", -1.77, 645, 667, -3.30, 0.81, "26.02.27", -5.84),
    Row("2026-03-26", 8, "HS2083", "恒生科技", "NO", -3.28, 4761, 4953, -3.88, 0.84, "26.03.19", -4.70),
    Row("2026-03-26", 9, "000510", "中证A500", "NO", -1.43, 5552, 5778, -3.91, 0.80, "26.03.13", -5.00),
    Row("2026-03-26", 10, "QQQ", "纳指100", "NO", -2.39, 573, 597, -4.02, 1.08, "26.03.12", -4.02),
    Row("2026-03-26", 11, "1B0016", "上证50", "NO", -1.22, 2824, 2945, -4.11, 0.75, "26.02.13", -8.34),
    Row("2026-03-26", 12, "932000", "中证2000", "NO", -1.58, 3234, 3422, -5.49, 0.94, "26.03.13", -6.83),
    Row("2026-03-26", 13, "1B0852", "中证1000", "NO", -1.44, 7639, 8079, -5.45, 0.89, "26.03.13", -7.00),
    Row("2026-03-26", 14, "399905", "中证500", "NO", -1.62, 7642, 8129, -5.99, 0.87, "26.03.12", -8.58),
    Row("2026-03-26", 15, "1B0688", "科创50", "NO", -2.02, 1288, 1371, -6.05, 0.69, "26.03.02", -12.02),
    Row("2026-03-26", 16, "899050", "北证50", "NO", -1.57, 1266, 1379, -8.19, 0.92, "26.02.27", -17.63, "现价从原文 12660 修正为 1266 以匹配偏离率。"),
    Row("2026-03-26", 17, "AUUSDO", "黄金现价", "NO", -2.03, 4404, 4931, -10.69, None, "26.03.12", -13.56),
    Row("2026-03-26", 18, "AGUSDO", "白银现价", "NO", -2.57, 68.82, 80.04, -14.02, None, "26.03.13", -14.57),
]

SECTOR_DATA: list[Row] = [
    Row("2026-03-31", 1, "399989", "中证医疗", "NO", -0.04, 6562, 6613, -0.77, 1.12, "26.01.27", -9.31),
    Row("2026-03-31", 2, "000922", "中证红利", "NO", -1.17, 5727, 5816, -1.53, 1.27, "26.03.18", -1.82),
    Row("2026-03-31", 3, "1B0932", "中证消费", "NO", -1.10, 14326, 14701, -2.55, 1.05, "26.03.19", -3.46),
    Row("2026-03-31", 4, "399941", "新能源", "NO", -3.56, 2708, 2803, -3.39, 0.83, "26.03.31", -0.44, "代码按 OCR 推断为 399941。"),
    Row("2026-03-31", 5, "399998", "中证煤炭", "NO", -3.96, 2342, 2442, -4.10, 1.04, "26.03.30", -3.98),
    Row("2026-03-31", 6, "000813", "细分化工", "NO", -2.75, 4013, 4194, -4.32, 0.84, "26.03.16", -6.57, "代码按 OCR 推断为 000813。"),
    Row("2026-03-31", 7, "886078", "商业航天", "NO", -0.99, 2249, 2367, -4.99, 1.08, "26.03.12", -9.28, "代码按 OCR 推断为 886078。"),
    Row("2026-03-31", 8, "931775", "房地产", "NO", -1.13, 2844, 2994, -4.99, 0.98, "26.02.26", -11.99, "代码按 OCR 推断，临界值按偏离率反推取整。"),
    Row("2026-03-31", 9, "399975", "证券公司", "NO", -0.80, 722, 764, -5.50, 0.99, "26.01.16", -14.96),
    Row("2026-03-31", 10, "1B0819", "有色金属", "NO", -1.47, 9616, 10191, -5.64, 0.98, "26.03.11", -12.84),
    Row("2026-03-31", 11, "881121", "半导体", "NO", -2.95, 1234, 1329, -7.18, 0.98, "26.03.03", -16.00, "现价/临界值/区间涨幅按 OCR 片段反推，需原图二次确认。"),
    Row("2026-03-31", 12, "H30590", "机器人", "NO", -1.03, 1721, 1832, -6.06, 0.94, "26.03.02", -13.52),
    Row("2026-03-31", 13, "881278", "电网设备", "NO", -1.48, 7192, 7664, -6.18, 1.06, "26.03.19", -4.07, "量比按 OCR 推断为 1.06。"),
    Row("2026-03-31", 14, "881279", "光伏设备", "NO", -3.31, 9151, 9804, -6.66, 0.76, "26.03.26", -5.26),
]


def build_markdown(rows: list[Row]) -> str:
    lines: list[str] = []
    lines.append("# 鱼盆趋势模型 v2.0 参考与校验报告")
    lines.append("")
    lines.append("## 结论")
    lines.append("")
    lines.append("- 用户提供的 2026-03-26、2026-03-30、2026-03-31 三期数据，绝大多数可以通过 `偏离率 = (现价 / 临界值点 - 1) × 100%` 精确复核。")
    lines.append("- 状态字段整体符合 `现价 < 临界值点 => NO`。")
    lines.append("- 当前项目代码 `web/src/utils/yupen.ts` 的 `calculateStrength` 按偏离率绝对值排序，这与用户提供的 v2.0 数据不一致；v2.0 更接近“偏离率越高、越接近临界值，趋势强度越强”。")
    lines.append("- 2026-03-30 的黄金行存在原始文本冲突，已在明细中保留备注，不作为模型错误。")
    lines.append("")
    lines.append("## 核心公式")
    lines.append("")
    lines.append("1. 偏离率 = (现价 / 临界值点 - 1) × 100%")
    lines.append("2. 状态 = YES 当且仅当现价 >= 临界值点，否则为 NO")
    lines.append("3. 区间涨幅 = (现价 / 状态转变日价格 - 1) × 100%")
    lines.append("4. 趋势强度 = 基于偏离率从高到低排序，而不是绝对值从大到小")
    lines.append("")
    for date in sorted({r.date for r in rows}, reverse=True):
        date_rows = [r for r in rows if r.date == date]
        valid_dev = [r for r in date_rows if r.deviation_match is True]
        invalid_dev = [r for r in date_rows if r.deviation_match is False]
        valid_status = [r for r in date_rows if r.status_match is True]
        invalid_status = [r for r in date_rows if r.status_match is False]
        lines.append(f"## {date} 明细")
        lines.append("")
        lines.append(f"- 偏离率校验通过: {len(valid_dev)}")
        lines.append(f"- 偏离率校验异常: {len(invalid_dev)}")
        lines.append(f"- 状态校验通过: {len(valid_status)}")
        lines.append(f"- 状态校验异常: {len(invalid_status)}")
        lines.append("")
        lines.append("| 强度 | 代码 | 名称 | 状态 | 涨幅% | 现价 | 临界值点 | 原始偏离率 | 复算偏离率 | 偏离率校验 | 状态校验 | 备注 |")
        lines.append("|---:|---|---|---|---:|---:|---:|---:|---:|---|---|---|")
        for r in sorted(date_rows, key=lambda x: x.strength):
            dev_ok = "通过" if r.deviation_match is True else "异常" if r.deviation_match is False else "跳过"
            status_ok = "通过" if r.status_match is True else "异常" if r.status_match is False else "跳过"
            lines.append(
                f"| {r.strength} | {r.code} | {r.name} | {r.status} | {fmt_pct(r.change_pct)} | "
                f"{fmt_number(r.current)} | {fmt_number(r.threshold)} | {fmt_pct(r.provided_deviation)} | "
                f"{fmt_pct(r.calculated_deviation)} | {dev_ok} | {status_ok} | {r.note or '-'} |"
            )
        lines.append("")
    lines.append("## 代码检查")
    lines.append("")
    lines.append("- `web/src/utils/yupen.ts` 中 `getDeviation()` 公式与用户数据一致。")
    lines.append("- `web/src/utils/yupen.ts` 中 `getStatus()` 采用 `>`，而参考数据更适合按 `>=` 解释。")
    lines.append("- `web/src/utils/yupen.ts` 中 `calculateStrength()` 当前实现为按偏离率绝对值降序，这会把深度负偏离资产排成“更强”，与 v2.0 排名方向相反。")
    lines.append("")
    return "\n".join(lines)


def build_docx(rows: list[Row]) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("鱼盆趋势模型 v2.0 参考与校验报告")
    set_cn_font(run, size=16)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("数据日期: 2026-03-26 / 2026-03-30 / 2026-03-31")
    set_cn_font(run, size=10)

    p = doc.add_paragraph()
    run = p.add_run(
        "说明: 本文档将用户提供的三期快照数据写入参考文档，并按偏离率公式与状态规则做复核。"
        "文本中存在 OCR 错位的行会保留备注，不将其误判为模型错误。"
    )
    set_cn_font(run)

    doc.add_heading("结论摘要", level=1)
    for text in [
        "核心公式可稳定反推出: 偏离率 = (现价 / 临界值点 - 1) × 100%。",
        "2026-03-31 数据内部高度自洽，SPY、HSI、QQQ、创业板指、沪深300、白银、北证50等行均可精确复算。",
        "当前项目代码的趋势强度排序逻辑与用户提供的 v2.0 排名不一致，需按偏离率从高到低重排，而不是按绝对值从大到小排序。",
        "2026-03-30 黄金行存在原始文本冲突，已标记为 OCR 疑点，不纳入状态是否匹配的结论。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        set_cn_font(run)

    doc.add_heading("核心公式", level=1)
    for text in [
        "偏离率 = (现价 / 临界值点 - 1) × 100%",
        "状态 = YES 当且仅当现价 >= 临界值点，否则为 NO",
        "区间涨幅 = (现价 / 状态转变日价格 - 1) × 100%",
        "趋势强度 = 基于偏离率从高到低排序，更接近临界值或已转正者更强",
    ]:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(text)
        set_cn_font(run)

    for date in sorted({r.date for r in rows}, reverse=True):
        date_rows = [r for r in rows if r.date == date]
        valid_dev = sum(r.deviation_match is True for r in date_rows)
        invalid_dev = sum(r.deviation_match is False for r in date_rows)
        valid_status = sum(r.status_match is True for r in date_rows)
        invalid_status = sum(r.status_match is False for r in date_rows)

        doc.add_heading(f"{date} 校验明细", level=1)
        p = doc.add_paragraph()
        run = p.add_run(
            f"偏离率通过 {valid_dev} 条, 异常 {invalid_dev} 条; "
            f"状态通过 {valid_status} 条, 异常 {invalid_status} 条。"
        )
        set_cn_font(run)

        table = doc.add_table(rows=1, cols=12)
        table.style = "Table Grid"
        headers = [
            "强度", "代码", "名称", "状态", "涨幅%", "现价", "临界值点",
            "原始偏离率", "复算偏离率", "偏离率校验", "状态校验", "备注"
        ]
        for idx, text in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = text
            shade_cell(cell, "D9EAF7")
            for para in cell.paragraphs:
                for run in para.runs:
                    set_cn_font(run)
                    run.bold = True

        for r in sorted(date_rows, key=lambda x: x.strength):
            row = table.add_row().cells
            values = [
                str(r.strength),
                r.code,
                r.name,
                r.status,
                fmt_pct(r.change_pct),
                fmt_number(r.current),
                fmt_number(r.threshold),
                fmt_pct(r.provided_deviation),
                fmt_pct(r.calculated_deviation),
                "通过" if r.deviation_match is True else "异常" if r.deviation_match is False else "跳过",
                "通过" if r.status_match is True else "异常" if r.status_match is False else "跳过",
                r.note or "-",
            ]
            for i, val in enumerate(values):
                row[i].text = val
                for para in row[i].paragraphs:
                    for run in para.runs:
                        set_cn_font(run, size=9)
                if i == 9:
                    if values[i] == "通过":
                        shade_cell(row[i], "E2F0D9")
                    elif values[i] == "异常":
                        shade_cell(row[i], "FCE4D6")
                if i == 10:
                    if values[i] == "通过":
                        shade_cell(row[i], "E2F0D9")
                    elif values[i] == "异常":
                        shade_cell(row[i], "FCE4D6")

    doc.add_heading("代码实现检查", level=1)
    for text in [
        "web/src/utils/yupen.ts 的 getDeviation 公式与参考数据一致。",
        "web/src/utils/yupen.ts 的 getStatus 当前采用 currentPrice > threshold，若要对齐参考报告，建议改为 >=。",
        "web/src/utils/yupen.ts 的 calculateStrength 当前按偏离率绝对值排序，这与 v2.0 参考数据不符，是最主要的逻辑差异。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        set_cn_font(run)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)


def main() -> None:
    OUTPUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_MD.write_text(build_markdown(DATA), encoding="utf-8")
    build_docx(DATA)
    print(OUTPUT_DOCX)
    print(OUTPUT_MD)


if __name__ == "__main__":
    main()
